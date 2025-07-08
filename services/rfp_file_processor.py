import os
import re
import logging
import tempfile
import pandas as pd
from typing import List, Dict, Any
from werkzeug.datastructures import FileStorage

logger = logging.getLogger(__name__)

class RFPFileProcessor:
    """Service for processing RFP files in various formats"""
    
    def __init__(self):
        self.supported_extensions = ['.xlsx', '.xls', '.csv', '.pdf', '.docx', '.txt']
    
    def process_rfp_file(self, file: FileStorage) -> Dict[str, Any]:
        """
        Process RFP file and extract line items
        
        Args:
            file: Uploaded RFP file
            
        Returns:
            Dict with extracted RFP items and metadata
        """
        try:
            # Validate file extension
            file_ext = os.path.splitext(file.filename.lower())[1]
            if file_ext not in self.supported_extensions:
                raise ValueError(f"Unsupported file format: {file_ext}")
            
            # Create temporary file for processing (Azure App Services compatible)
            file_extension = os.path.splitext(file.filename)[1]
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=file_extension)
            temp_file_path = temp_file.name

            # Save uploaded file to temporary location
            file.save(temp_file_path)
            temp_file.close()
            
            logger.info(f"Processing RFP file: {file.filename}, size: {os.path.getsize(temp_file_path)} bytes")
            
            # Process based on file type
            if file_ext in ['.xlsx', '.xls']:
                rfp_items = self._process_excel_rfp(temp_file_path)
            elif file_ext == '.csv':
                rfp_items = self._process_csv_rfp(temp_file_path)
            elif file_ext == '.pdf':
                rfp_items = self._process_pdf_rfp(temp_file_path)
            elif file_ext == '.docx':
                rfp_items = self._process_docx_rfp(temp_file_path)
            elif file_ext == '.txt':
                rfp_items = self._process_text_rfp(temp_file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
            
            # Clean up temporary file
            os.remove(temp_file_path)
            
            # Clean and filter RFP items
            cleaned_items = self._clean_rfp_items(rfp_items)
            
            logger.info(f"Extracted {len(cleaned_items)} RFP items from {file.filename}")
            
            return {
                'success': True,
                'total_items': len(cleaned_items),
                'rfp_items': cleaned_items,
                'file_type': file_ext,
                'original_filename': file.filename
            }
            
        except Exception as e:
            logger.error(f"Error processing RFP file: {str(e)}")
            # Clean up temporary file if it exists
            if 'temp_file_path' in locals() and os.path.exists(temp_file_path):
                os.remove(temp_file_path)
            raise e
    
    def _process_excel_rfp(self, file_path: str) -> List[str]:
        """Process Excel RFP file"""
        try:
            # Try to read Excel file
            df = pd.read_excel(file_path, engine='openpyxl')
            
            # Look for columns that might contain RFP items
            potential_columns = []
            for col in df.columns:
                col_lower = str(col).lower()
                if any(keyword in col_lower for keyword in ['description', 'item', 'material', 'spec', 'product', 'part']):
                    potential_columns.append(col)
            
            # If no obvious columns found, use all text columns
            if not potential_columns:
                potential_columns = [col for col in df.columns if df[col].dtype == 'object']
            
            # Extract text from potential columns
            rfp_items = []
            for col in potential_columns:
                for value in df[col].dropna():
                    if isinstance(value, str) and len(value.strip()) > 10:
                        rfp_items.append(value.strip())
            
            return rfp_items
            
        except Exception as e:
            logger.error(f"Error processing Excel RFP: {e}")
            return []
    
    def _process_csv_rfp(self, file_path: str) -> List[str]:
        """Process CSV RFP file"""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            # Look for columns that might contain RFP items
            potential_columns = []
            for col in df.columns:
                col_lower = str(col).lower()
                if any(keyword in col_lower for keyword in ['description', 'item', 'material', 'spec', 'product', 'part']):
                    potential_columns.append(col)
            
            # If no obvious columns found, use all text columns
            if not potential_columns:
                potential_columns = [col for col in df.columns if df[col].dtype == 'object']
            
            # Extract text from potential columns
            rfp_items = []
            for col in potential_columns:
                for value in df[col].dropna():
                    if isinstance(value, str) and len(value.strip()) > 10:
                        rfp_items.append(value.strip())
            
            return rfp_items
            
        except Exception as e:
            logger.error(f"Error processing CSV RFP: {e}")
            return []
    
    def _process_pdf_rfp(self, file_path: str) -> List[str]:
        """Process PDF RFP file"""
        try:
            # Try pdfplumber first (better for tables)
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    text = ""
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    
                    if text.strip():
                        return self._extract_items_from_text(text)
            except ImportError:
                logger.warning("pdfplumber not available, trying PyPDF2")
            
            # Fallback to PyPDF2
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                    
                    return self._extract_items_from_text(text)
            except ImportError:
                logger.error("No PDF processing library available")
                return []
            
        except Exception as e:
            logger.error(f"Error processing PDF RFP: {e}")
            return []
    
    def _process_docx_rfp(self, file_path: str) -> List[str]:
        """Process DOCX RFP file"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
            
            return self._extract_items_from_text(text)
            
        except ImportError:
            logger.error("python-docx not available")
            return []
        except Exception as e:
            logger.error(f"Error processing DOCX RFP: {e}")
            return []
    
    def _process_text_rfp(self, file_path: str) -> List[str]:
        """Process plain text RFP file"""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            return self._extract_items_from_text(text)
            
        except Exception as e:
            logger.error(f"Error processing text RFP: {e}")
            return []
    
    def _extract_items_from_text(self, text: str) -> List[str]:
        """Extract RFP items from raw text"""
        lines = text.split('\n')
        items = []
        
        for line in lines:
            line = line.strip()
            
            # Skip empty lines and headers
            if not line or len(line) < 10:
                continue
            
            # Skip lines that look like headers or page numbers
            if any(keyword in line.lower() for keyword in ['page', 'total', 'subtotal', 'header', 'footer']):
                continue
            
            # Look for lines that contain technical specifications
            if self._looks_like_rfp_item(line):
                items.append(line)
        
        return items
    
    def _looks_like_rfp_item(self, line: str) -> bool:
        """Check if a line looks like an RFP item"""
        line_lower = line.lower()
        
        # Technical keywords that suggest this is an RFP item
        technical_keywords = [
            'pipe', 'flange', 'elbow', 'tee', 'reducer', 'valve', 'fitting',
            'bolt', 'nut', 'gasket', 'coupling', 'nipple', 'cap', 'plug',
            'sch', 'schedule', 'npt', 'bw', 'sw', 'rf', 'ff', 'wnrf',
            'a105', 'a234', 'a106', 'ss316', 'ss304', 'carbon steel',
            'stainless', 'api', 'astm', 'asme', 'ansi', 'din', 'iso'
        ]
        
        # Size patterns
        size_patterns = [
            r'\d+["\']',  # 6", 4'
            r'\d+\.\d+["\']',  # 6.5"
            r'\d+mm',  # 150mm
            r'dn\d+',  # DN100
            r'nb\d+',  # NB150
            r'\d+#',   # 150#
        ]
        
        # Check for technical keywords
        has_technical_keyword = any(keyword in line_lower for keyword in technical_keywords)
        
        # Check for size patterns
        has_size_pattern = any(re.search(pattern, line_lower) for pattern in size_patterns)
        
        # Must have either technical keyword or size pattern, and be reasonably long
        return (has_technical_keyword or has_size_pattern) and len(line) > 15
    
    def _clean_rfp_items(self, items: List[str]) -> List[str]:
        """Clean and deduplicate RFP items"""
        cleaned = []
        seen = set()
        
        for item in items:
            # Clean the item
            cleaned_item = re.sub(r'\s+', ' ', item.strip())
            
            # Remove duplicates (case insensitive)
            if cleaned_item.lower() not in seen and len(cleaned_item) > 10:
                cleaned.append(cleaned_item)
                seen.add(cleaned_item.lower())
        
        return cleaned
